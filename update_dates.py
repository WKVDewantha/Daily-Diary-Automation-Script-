from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from datetime import datetime, timedelta
import re
import os

# ==========================================
# Settings
# ==========================================
input_file = "Daily Dairy uvt ict.docx"        # Name of your original file
output_file = "Daily Dairy Final by A_L_E_X.docx" # Name of the new output file

# Start Date: November 20, 2025 (Thursday)
start_date_user = datetime(2025, 11, 20) 
# Calculate the Monday of that week to align with the diary format
start_monday = start_date_user - timedelta(days=start_date_user.weekday()) # Nov 17

total_weeks = 26
day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']

# Holiday List
holiday_map = {
    (2025, 11, 28): "Bad Weather Holiday",
    (2025, 12, 4): "Unduvap Full Moon Poya Day",
    (2025, 12, 25): "Christmas Day",
    (2025, 12, 26): "Government Holiday",
    (2026, 1, 3): "Duruthu Full Moon Poya Day",
    (2026, 1, 15): "Tamil Thai Pongal Day",
    (2026, 2, 1): "Navam Full Moon Poya Day",
    (2026, 2, 4): "National Day",
    (2026, 2, 15): "Mahasivarathri Day",
    (2026, 3, 2): "Madin Full Moon Poya Day",
    (2026, 3, 21): "Id Ul-Fitr",
    (2026, 4, 1): "Bak Full Moon Poya Day",
    (2026, 4, 3): "Good Friday",
    (2026, 4, 13): "Day prior to Sinhala & Tamil New Year Day",
    (2026, 4, 14): "Sinhala & Tamil New Year Day",
    (2026, 5, 1): "May Day / Vesak Full Moon Poya Day",
    (2026, 5, 2): "Day following Vesak Full Moon Poya Day"
}

# Date Generation
weeks_data = []
for i in range(total_weeks):
    week_start = start_monday + timedelta(weeks=i)
    week_dates = []
    for d in range(7):
        week_dates.append(week_start + timedelta(days=d))
    weeks_data.append(week_dates)

# ==========================================
# Helper Functions
# ==========================================
def safe_replace_date_in_cell(cell, new_date_text):
    """Updates the date in a cell without changing the existing font style."""
    if not cell.paragraphs:
        cell.add_paragraph(new_date_text)
        return
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        # Replace the text of the first run to preserve formatting
        paragraph.runs[0].text = new_date_text
        # Clear remaining runs to avoid duplication
        for i in range(1, len(paragraph.runs)):
            paragraph.runs[i].text = ""
    else:
        paragraph.add_run(new_date_text)

def safe_replace_header_date(cell, sunday_date_str):
    """Updates the 'Week Ending' Sunday date in the header without changing font."""
    for paragraph in cell.paragraphs:
        if "Sunday" in paragraph.text:
            found = False
            for run in paragraph.runs:
                if "Sunday" in run.text or re.search(r'\d{2}/\d{2}/\d{2}', run.text):
                    # Use Regex to find 'Sunday' followed by any date pattern
                    new_text = re.sub(r'Sunday\s*[\d/]*', f'Sunday {sunday_date_str}', run.text)
                    if new_text == run.text and "Sunday" in run.text:
                         run.text = f"Sunday {sunday_date_str}"
                    else:
                        run.text = new_text
                    found = True
                    break
            if found: break

# ==========================================
# Main Logic
# ==========================================
try:
    print("🔄 Starting document update process...")
    doc = Document(input_file)
    week_counter = 0

    for table in doc.tables:
        if week_counter >= total_weeks:
            break

        # --- 1. Header Date Update ---
        is_header = False
        for row in table.rows:
            for cell in row.cells:
                if "FOR THE WEEK ENDING" in cell.text:
                    is_header = True
                    sunday_date = weeks_data[week_counter][6].strftime("%y/%m/%d")
                    safe_replace_header_date(cell, sunday_date)
                    break
            if is_header: break

        # --- 2. Schedule Rows Update ---
        day_row_map = {}
        for r_idx, row in enumerate(table.rows):
            row_txt = " ".join([c.text for c in row.cells])
            for d_idx, day in enumerate(day_names):
                if day in row_txt:
                    day_row_map[d_idx] = r_idx

        if day_row_map:
            print(f"   - Updating Week {week_counter + 1}...")
            current_week = weeks_data[week_counter]

            for day_idx, date_obj in enumerate(current_week):
                if day_idx in day_row_map:
                    row = table.rows[day_row_map[day_idx]]
                    date_str = date_obj.strftime("%y/%m/%d")

                    # (A) Update Date (Column 1)
                    if len(row.cells) > 1:
                        safe_replace_date_in_cell(row.cells[1], date_str)

                    # (B) Update Description (Column 2) - Logic for Weekends & Holidays
                    if len(row.cells) > 2:
                        cell = row.cells[2]
                        date_key = (date_obj.year, date_obj.month, date_obj.day)
                        
                        # LOGIC PRIORITY:
                        # 1. Check if it is Saturday (5) or Sunday (6). If yes, mark as WEEKEND.
                        if day_idx in [5, 6]: 
                            # --- WEEKEND (Black Color) ---
                            cell.text = "" 
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
                            p = cell.paragraphs[0]
                            r = p.add_run("WEEKEND")
                            r.bold = True
                            r.font.color.rgb = RGBColor(127, 127, 127) # GRAY
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 2. If it is NOT a weekend, check if it is a Holiday.
                        elif date_key in holiday_map:
                            # --- HOLIDAY (Red Color) ---
                            cell.text = "" 
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
                            p = cell.paragraphs[0]
                            r = p.add_run(holiday_map[date_key])
                            r.bold = True
                            r.font.color.rgb = RGBColor(255, 0, 0) # RED
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            week_counter += 1

    doc.save(output_file)
    print(f"\n✅ Suprise! Lazy M#F... Ur document has been successfully updated.")
    print(f"📄 Saved as: {output_file}")

except Exception as e:
    print(f"\n❌ Error: {e}")

# ==========================================
# Keep Window Open
# ==========================================
print("\n" + "="*30)
input("Press Enter to go away...")
